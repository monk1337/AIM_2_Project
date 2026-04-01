"""Generate the EgoSurg-Bench data collection plan as a .docx file."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pathlib import Path


def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elem)


def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    return table


def main():
    doc = Document()

    # --- Styles ---
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # --- Title ---
    title = doc.add_heading('EgoSurg-Bench: Data Collection Plan', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Egocentric Open Surgery Training Dataset for VLA Fine-Tuning')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('Project: Surgical VLA (VITRA Fine-Tuning)\n').bold = True
    info.add_run('Prepared for: Dr. Brat, Guillaume Kugener MD\n')
    info.add_run('Author: Aaditya (Harvard BMI 702)\n')
    info.add_run('Date: March 2026')

    doc.add_page_break()

    # --- Table of Contents ---
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Executive Summary',
        '2. Motivation & Gap Analysis',
        '3. Camera Setup & Calibration',
        '4. Task Inventory (20 Tasks)',
        '5. Recording Protocol',
        '6. Metadata Schema',
        '7. During-Recording Annotations',
        '8. Post-Recording Processing Pipeline',
        '9. Volume & Scale Estimates',
        '10. Recording Schedule',
        '11. Equipment Checklist',
        '12. Future-Proofing',
        '13. Comparison with Existing Datasets',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # =================================================================
    # 1. EXECUTIVE SUMMARY
    # =================================================================
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'We propose recording an egocentric surgical training dataset ("EgoSurg-Bench") '
        'on synthetic training pads with real instruments, real gloved hands, and real '
        'environment conditions. This dataset will be the first to combine:'
    )
    bullets = [
        'High frame-rate egocentric video (30fps) enabling temporal hand tracking via HaWoR',
        'Calibrated camera intrinsics (exact K matrix, not estimated)',
        'Multiple surgical task types (20 tasks across 5 categories)',
        'Multiple operators at different skill levels (enabling skill assessment research)',
        'Audio narration for language annotation generation',
        'JIGSAWS-comparable protocol for direct benchmarking',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_paragraph(
        '\nThe recorded data will be processed through the VITRA pipeline (HaWoR hand '
        'reconstruction + MegaSAM camera poses + VLM language labeling) to produce '
        'VITRA-compatible .npy episodes for fine-tuning the VITRA foundation model on '
        'surgical manipulation.'
    )

    # =================================================================
    # 2. MOTIVATION & GAP ANALYSIS
    # =================================================================
    doc.add_heading('2. Motivation & Gap Analysis', level=1)

    doc.add_heading('Why we need a new dataset', level=2)
    doc.add_paragraph(
        'No existing public dataset provides all components needed for VITRA fine-tuning '
        'on open surgery. The table below shows what each dataset provides vs. what VITRA requires:'
    )

    add_table(doc,
        ['Component', 'EgoSurgery', 'AVOS', 'JIGSAWS', 'POV-Surgery', 'VITRA needs'],
        [
            ['Egocentric video', 'JPGs at 0.5fps', 'No (3rd person)', 'No (overhead)', 'Yes (synthetic)', 'Yes, 30fps'],
            ['Hand pose (MANO)', 'None', 'None', 'Robot kinematics', 'Yes (GT)', 'Yes'],
            ['Camera intrinsics', 'Unknown', 'Unknown', 'N/A', 'Known', 'Yes (exact)'],
            ['Camera extrinsics', 'None', 'None', 'N/A', 'Known', 'Yes (per-frame)'],
            ['Language annotations', 'None', 'Action classes', 'Gesture labels', 'None', 'Yes (per-hand text)'],
            ['Temporal sequences', '0.5fps frames', 'Video', 'Video', 'Frames', 'Video at >= 25fps'],
            ['Open surgery', 'Yes', 'Yes', 'No (dVRK)', 'Yes', 'Our target domain'],
            ['Multiple operators', '8 surgeons', 'Unknown', '8 users', 'N/A', 'Yes (skill eval)'],
        ]
    )

    doc.add_paragraph()
    doc.add_paragraph(
        'Key limitation of our current pipeline: EgoSurgery provides only 0.5fps JPG frames, '
        'which is too sparse for HaWoR (requires video for temporal hand tracking and SLAM). '
        'WiLoR (single-frame estimator) produces noisy hand poses on surgical gloved hands. '
        'Recording our own high-fps video solves both problems.'
    )

    # =================================================================
    # 3. CAMERA SETUP & CALIBRATION
    # =================================================================
    doc.add_heading('3. Camera Setup & Calibration', level=1)

    doc.add_heading('Camera configuration', level=2)
    add_table(doc,
        ['Camera', 'Purpose', 'Specs', 'Priority'],
        [
            ['Head-mounted (ego)', 'Primary VITRA input', '1080p, 30fps min, wide-angle', 'Required'],
            ['Overhead fixed (exo)', 'Multi-view training (EgoExo4D-style)', '1080p, 30fps', 'Recommended'],
            ['Depth camera', '3D ground truth validation', 'Intel RealSense D435', 'Optional'],
        ]
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Camera mounting: ').bold = True
    p.add_run(
        'Head-mounted camera must be on the forehead, angled ~30-45 degrees downward '
        'toward the hands. Chest-mounted causes too much hand occlusion. '
        'GoPro Hero with head strap mount or Tobii Pro Glasses are recommended options.'
    )

    doc.add_heading('Calibration procedure (one-time, ~15 minutes)', level=2)
    add_table(doc,
        ['What', 'How', 'Why'],
        [
            ['Camera intrinsics', 'Checkerboard pattern (A3 print), ~20 photos, OpenCV calibrateCamera()', 'Exact K matrix for HaWoR and VITRA'],
            ['Stereo extrinsics', 'Checkerboard visible to both ego+exo simultaneously', 'Multi-view reconstruction'],
            ['Color calibration', 'X-Rite ColorChecker card, one photo per session', 'Consistent appearance across sessions'],
            ['Workspace scale', 'Measure training pad dimensions, place ruler in frame', 'Metric-scale world-space reference'],
        ]
    )

    # =================================================================
    # 4. TASK INVENTORY
    # =================================================================
    doc.add_heading('4. Task Inventory (20 Tasks)', level=1)
    doc.add_paragraph(
        'Tasks are organized into 5 categories covering all fundamental open surgical '
        'hand skills. Each task produces distinct hand pose distributions, ensuring the '
        'model learns diverse manipulation patterns.'
    )

    # Category 1: Suturing
    doc.add_heading('Category 1: Suturing (7 tasks)', level=2)
    add_table(doc,
        ['#', 'Task', 'Instruments', 'Duration', 'Distinct hand motion'],
        [
            ['1', 'Simple interrupted suture', 'Needle driver + forceps', '3 min', 'Needle driving + knot per stitch'],
            ['2', 'Running (continuous) suture', 'Needle driver + forceps', '3 min', 'Repeated passes, no inter-stitch knots'],
            ['3', 'Horizontal mattress suture', 'Needle driver + forceps', '3 min', 'Double pass, wider tissue bite'],
            ['4', 'Vertical mattress suture', 'Needle driver + forceps', '3 min', 'Deep-shallow-shallow-deep pattern'],
            ['5', 'Subcuticular suture', 'Needle driver + forceps', '3 min', 'Intradermal horizontal bites'],
            ['6', 'Figure-of-eight suture', 'Needle driver + forceps', '2 min', 'Crossing pattern'],
            ['7', 'Needle loading + transfer', '2 needle drivers', '2 min', 'Precision grip, hand-to-hand pass'],
        ]
    )

    # Category 2: Knot Tying
    doc.add_heading('Category 2: Knot Tying (4 tasks)', level=2)
    add_table(doc,
        ['#', 'Task', 'Instruments', 'Duration', 'Distinct hand motion'],
        [
            ['8', 'Two-handed instrument tie', 'Needle driver + forceps', '2 min', 'Wrap-grab-pull alternating hands'],
            ['9', 'One-handed instrument tie', 'Needle driver only', '2 min', 'Single hand wrapping'],
            ['10', 'Hand tie (square knot)', 'Bare hands + suture', '2 min', 'Finger loops, no instruments'],
            ['11', "Surgeon's knot (double throw)", 'Bare hands or instruments', '2 min', 'Double first throw for tension'],
        ]
    )

    # Category 3: Cutting & Incision
    doc.add_heading('Category 3: Cutting & Incision (4 tasks)', level=2)
    add_table(doc,
        ['#', 'Task', 'Instruments', 'Duration', 'Distinct hand motion'],
        [
            ['12', 'Straight incision with scalpel', 'Scalpel + forceps', '1 min', 'Drawing cut, stabilizing with forceps'],
            ['13', 'Curved incision with scalpel', 'Scalpel + forceps', '1 min', 'Wrist rotation during cut'],
            ['14', 'Scissors cutting (tissue)', 'Metzenbaum scissors + forceps', '1 min', 'Open-advance-close pattern'],
            ['15', 'Suture cutting', 'Suture scissors', '1 min', 'Quick snip at specific length'],
        ]
    )

    # Category 4: Dissection & Tissue Handling
    doc.add_heading('Category 4: Dissection & Tissue Handling (3 tasks)', level=2)
    add_table(doc,
        ['#', 'Task', 'Instruments', 'Duration', 'Distinct hand motion'],
        [
            ['16', 'Sharp dissection', 'Metzenbaum scissors + forceps', '2 min', 'Spread-cut along plane'],
            ['17', 'Blunt dissection', 'Two forceps', '2 min', 'Spread-separate motion'],
            ['18', 'Tissue retraction + exposure', 'Forceps + retractor/fingers', '2 min', 'Hold tension while other hand works'],
        ]
    )

    # Category 5: Combined / Procedural
    doc.add_heading('Category 5: Combined / Procedural (2 tasks)', level=2)
    add_table(doc,
        ['#', 'Task', 'Instruments', 'Duration', 'Distinct hand motion'],
        [
            ['19', 'Full wound closure (incision > suture > knot > cut)', 'All', '5 min', 'Chained multi-step procedure'],
            ['20', 'Debridement simulation', 'Forceps + scissors', '2 min', 'Grip-lift-cut irregular pattern'],
        ]
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Why 20 tasks: ').bold = True
    p.add_run(
        'Covers all fundamental open surgical hand skills. Any open procedure is a '
        'combination of these primitives. Each task produces a distinct hand pose '
        'distribution (scalpel grip vs needle driver grip vs forceps grip vs bare hand). '
        'Task 19 (full wound closure) tests multi-step action sequencing. '
        'Tasks 1, 7, 8, 10 directly map to JIGSAWS tasks for benchmarking.'
    )

    # =================================================================
    # 5. RECORDING PROTOCOL
    # =================================================================
    doc.add_heading('5. Recording Protocol', level=1)

    doc.add_heading('Per-surgeon session flow', level=2)
    steps = [
        ('Setup (15 min)', 'Mount cameras, calibrate if first session, set up training pad and instruments, verify recording quality with test clip.'),
        ('Warm-up (5 min)', 'Surgeon performs a few practice stitches (not recorded). Adjusts camera angle if needed.'),
        ('Recording blocks (3-3.5 hrs)', 'Work through all 20 tasks, 5 trials each. Take 5-minute breaks every 30 minutes. Announce task name and trial number at the start of each recording.'),
        ('Wrap-up (10 min)', 'Verify all recordings saved. Note any issues. Quick debrief.'),
    ]
    for title_text, desc in steps:
        p = doc.add_paragraph()
        p.add_run(f'{title_text}: ').bold = True
        p.add_run(desc)

    doc.add_heading('Per-recording procedure', level=2)
    recording_steps = [
        'Start recording on all cameras simultaneously',
        'Surgeon announces: "[Task name], trial [N], starting now"',
        'Surgeon performs the task while narrating actions aloud',
        'Surgeon announces "Done" at completion',
        'Stop recording',
        'Brief pause before next trial (~15 seconds)',
    ]
    for i, step in enumerate(recording_steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_heading('Participant requirements', level=2)
    add_table(doc,
        ['Parameter', 'Minimum', 'Ideal', 'Why'],
        [
            ['Number of surgeons', '3', '5-8', 'Leave-one-out evaluation requires >= 3; JIGSAWS uses 8'],
            ['Skill levels', '1 level', 'Mix: PGY-1, PGY-3, PGY-5, Attending', 'Skill assessment research + diversity'],
            ['Trials per task', '3', '5', 'Statistical reliability + learning curve analysis'],
            ['Dominant hand', 'Note it', 'Include both L and R dominant', 'Hand-agnostic model training'],
        ]
    )

    # =================================================================
    # 6. METADATA SCHEMA
    # =================================================================
    doc.add_heading('6. Metadata Schema', level=1)
    doc.add_paragraph(
        'Each recording gets a structured metadata JSON file. This metadata enables '
        'filtering, benchmarking, and future extensions without re-processing video.'
    )

    add_table(doc,
        ['Field', 'Type', 'Example', 'Notes'],
        [
            ['recording_id', 'string', 'S01_T03_trial02', 'Surgeon_Task_Trial'],
            ['surgeon_id', 'string', 'S01', 'Anonymized'],
            ['skill_level', 'enum', 'intermediate', 'novice / intermediate / expert'],
            ['years_experience', 'int', '4', 'Post-residency years'],
            ['dominant_hand', 'string', 'right', 'left / right'],
            ['task_id', 'int', '3', '1-20 from task inventory'],
            ['task_name', 'string', 'horizontal_mattress_suture', 'Snake_case'],
            ['task_category', 'string', 'suturing', '5 categories'],
            ['trial_number', 'int', '2', '1-5'],
            ['instruments', 'list', '["needle_driver", "forceps"]', 'All instruments used'],
            ['suture_type', 'string', '4-0 Vicryl', 'Material and size'],
            ['pad_type', 'string', 'Ethicon LPSE-1', 'Training pad model'],
            ['glove_type', 'string', 'Ansell Micro-Touch, size 7', 'Brand, material, size'],
            ['glove_color', 'string', 'white', 'Affects hand tracking'],
            ['camera_ego', 'string', 'GoPro Hero 12', 'Model'],
            ['camera_exo', 'string', 'Canon EOS R50', 'Model (if used)'],
            ['intrinsics_file', 'path', 'calibration/gopro12.json', 'Pre-computed K matrix'],
            ['fps', 'int', '30', 'Frame rate'],
            ['resolution', 'string', '1920x1080', 'Video resolution'],
            ['date', 'string', '2026-03-15', 'ISO date'],
            ['location', 'string', 'Lahey Hospital sim lab', 'Recording location'],
            ['duration_seconds', 'float', '185.3', 'Actual recording length'],
            ['has_audio_narration', 'bool', 'true', 'Whether surgeon narrated'],
            ['has_exo_view', 'bool', 'true', 'Whether overhead camera was used'],
            ['has_depth', 'bool', 'false', 'Whether depth camera was used'],
            ['notes', 'string', 'Needle dropped at 1:23', 'Free-form notes'],
        ]
    )

    # =================================================================
    # 7. DURING-RECORDING ANNOTATIONS
    # =================================================================
    doc.add_heading('7. During-Recording Annotations', level=1)
    doc.add_paragraph(
        'These annotations are collected live during recording, requiring minimal extra '
        'effort from the surgeon but providing high-value ground truth.'
    )

    add_table(doc,
        ['Annotation', 'How collected', 'Format', 'Value'],
        [
            ['Audio narration', 'Surgeon speaks aloud: "I\'m driving the needle through the far edge"', 'Audio track on video', 'Transcribed via Whisper -> language annotations for VITRA'],
            ['Action boundaries', 'Verbal cue: "starting suture", "tying knot now", "done"', 'Timestamps from audio', 'Atomic action segmentation ground truth'],
            ['Error flagging', 'Surgeon says "error" or "mistake" when they make one', 'Timestamps', 'Skill assessment, error detection research'],
            ['Instrument changes', 'Surgeon announces when switching instruments', 'Timestamps', 'Tool usage tracking, action triplets'],
            ['Difficulty rating', 'Surgeon rates 1-5 after each trial', 'Post-trial verbal', 'Subjective difficulty for skill modeling'],
        ]
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Important: ').bold = True
    p.add_run(
        'Narration should be natural and brief, not scripted. The goal is to capture '
        'what the surgeon would say when teaching a resident ("now I\'m loading the needle '
        'backhand because of the angle") rather than formal descriptions. This produces '
        'more realistic and diverse language annotations.'
    )

    # =================================================================
    # 8. POST-RECORDING PROCESSING
    # =================================================================
    doc.add_heading('8. Post-Recording Processing Pipeline', level=1)
    doc.add_paragraph(
        'After recording, the following processing pipeline converts raw video into '
        'VITRA-compatible training episodes:'
    )

    add_table(doc,
        ['Step', 'Tool', 'Input', 'Output'],
        [
            ['1. Sync cameras', 'Clap/flash alignment', 'Ego + exo video', 'Time-aligned multi-view'],
            ['2. Camera intrinsics', 'OpenCV (pre-computed)', 'Calibration images', 'K matrix (3x3)'],
            ['3. Camera poses', 'MegaSAM + MoGe-2', 'Ego video', 'Per-frame extrinsics (T, 4, 4)'],
            ['4. Hand reconstruction', 'HaWoR', 'Ego video', 'World-space MANO per frame (61-dim per hand)'],
            ['5. Temporal smoothing', 'Spline fitting', 'Raw MANO trajectories', 'Smooth hand trajectories'],
            ['6. Transcribe narration', 'OpenAI Whisper', 'Audio track', 'Timestamped text transcript'],
            ['7. Language annotations', 'VLM (GPT/Claude) on transcript + frames', 'Transcript + key frames', 'Per-hand action instructions'],
            ['8. Action segmentation', 'Speed minima detection (VITRA method)', '3D hand trajectories', 'Atomic action boundaries'],
            ['9. Episode packaging', 'Custom script (step3 adapted)', 'All above', 'VITRA .npy episodes'],
            ['10. Skill scoring', 'OSATS rubric by surgeon reviewer', 'Video recordings', 'Ground-truth skill labels per trial'],
        ]
    )

    # =================================================================
    # 9. VOLUME & SCALE ESTIMATES
    # =================================================================
    doc.add_heading('9. Volume & Scale Estimates', level=1)

    add_table(doc,
        ['Parameter', 'Conservative (3 surgeons)', 'Good (5 surgeons)', 'Ideal (8 surgeons)'],
        [
            ['Tasks', '20', '20', '20'],
            ['Trials per task', '3', '5', '5'],
            ['Total recordings', '180', '500', '800'],
            ['Total recording time', '~6 hrs', '~17 hrs', '~28 hrs'],
            ['Total frames (30fps)', '~6.5M', '~18M', '~30M'],
            ['Atomic episodes (est.)', '~2,000', '~5,000', '~8,000'],
            ['Storage (raw video)', '~50 GB', '~140 GB', '~230 GB'],
        ]
    )

    doc.add_paragraph()
    doc.add_heading('Comparison with existing datasets', level=2)
    add_table(doc,
        ['Dataset', 'Recordings', 'Tasks', 'Users', 'Domain', 'Hand pose'],
        [
            ['JIGSAWS', '101', '3', '8', 'Laparoscopic (dVRK)', 'Robot kinematics'],
            ['SutureBot', '1,890', '1 (suturing)', 'Multiple', 'Laparoscopic (dVRK)', 'Robot kinematics'],
            ['EgoSurgery', '~20 videos', '9 phases', '8', 'Open (real)', 'None'],
            ['POV-Surgery', '88K frames', 'Multiple', 'Synthetic', 'Open (synthetic)', 'MANO GT'],
            ['EgoSurg-Bench (ours, ideal)', '800', '20', '8', 'Open (training pad)', 'HaWoR MANO'],
        ]
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Fine-tuning math: ').bold = True
    p.add_run(
        'VITRA pretraining used 1.22M episodes. For domain-specific fine-tuning, '
        '2,000-8,000 episodes of high-quality, in-domain data is sufficient '
        '(analogous to how VITRA fine-tunes on ~5K robot episodes for real-world deployment). '
        'Quality and domain relevance matter more than quantity at this stage.'
    )

    # =================================================================
    # 10. RECORDING SCHEDULE
    # =================================================================
    doc.add_heading('10. Recording Schedule', level=1)

    add_table(doc,
        ['Scenario', 'Per surgeon', 'Surgeons', 'Total time', 'Calendar'],
        [
            ['Minimum viable', '60 recordings (~2 hrs)', '3', '~6 hrs', '3 half-day sessions'],
            ['Good', '100 recordings (~3.5 hrs)', '5', '~17 hrs', '5 sessions over 2-3 weeks'],
            ['Ideal', '100 recordings (~3.5 hrs)', '8', '~28 hrs', '8 sessions over 1 month'],
        ]
    )

    doc.add_paragraph()
    doc.add_paragraph(
        'Each session structure: 15 min setup + 5 min warm-up + 3-3.5 hrs recording '
        '(with breaks) + 10 min wrap-up = ~4 hours total per surgeon.'
    )

    doc.add_heading('Suggested session breakdown per surgeon', level=2)
    add_table(doc,
        ['Block', 'Tasks', 'Trials', 'Time'],
        [
            ['Block 1', 'Tasks 1-5 (Suturing)', '5 each', '~60 min'],
            ['Break', '', '', '10 min'],
            ['Block 2', 'Tasks 6-11 (Suturing + Knots)', '5 each', '~50 min'],
            ['Break', '', '', '10 min'],
            ['Block 3', 'Tasks 12-18 (Cutting + Dissection)', '5 each', '~50 min'],
            ['Break', '', '', '10 min'],
            ['Block 4', 'Tasks 19-20 (Combined)', '5 each', '~30 min'],
        ]
    )

    # =================================================================
    # 11. EQUIPMENT CHECKLIST
    # =================================================================
    doc.add_heading('11. Equipment Checklist', level=1)

    doc.add_heading('Cameras & Electronics', level=2)
    items = [
        'Head-mounted camera (GoPro Hero 12 + head strap mount)',
        'Overhead tripod camera (any 1080p/30fps)',
        'Micro-SD cards (128GB minimum per camera per session)',
        'Spare batteries for GoPro',
        'Lapel microphone (or rely on GoPro built-in audio)',
        'Laptop for verifying recordings between blocks',
        'Checkerboard calibration board (A3 printed, rigid backing)',
        'Color calibration card (X-Rite ColorChecker, optional)',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Surgical Instruments', level=2)
    items = [
        'Needle drivers (x2) — e.g., Mayo-Hegar',
        'Tissue forceps (Adson with teeth)',
        'Metzenbaum scissors',
        'Suture scissors',
        'Scalpel handle (#3) + blades (#10, #15)',
        'Retractor (small Army-Navy or Senn)',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Consumables', level=2)
    items = [
        'Surgical training pads (e.g., Ethicon LPSE-1 or similar) — 2-3 per session',
        'Suture material: 3-0 and 4-0 Vicryl (or Prolene) on curved needle',
        'Surgical gloves — white, known brand/size (Ansell Micro-Touch recommended)',
        'Extra gloves (multiple sizes for different surgeons)',
        'Drapes or towels for workspace setup',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Environment', level=2)
    items = [
        'Well-lit workspace (overhead surgical lamp or ring light)',
        'Stable table/surface for training pad',
        'Minimal background clutter (consistent background helps visual models)',
        'Room with minimal interruptions',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    # =================================================================
    # 12. FUTURE-PROOFING
    # =================================================================
    doc.add_heading('12. Future-Proofing', level=1)
    doc.add_paragraph(
        'This dataset is designed to be useful beyond our immediate VITRA fine-tuning goal. '
        'The following design decisions ensure long-term value:'
    )

    add_table(doc,
        ['Design decision', 'Immediate value', 'Future value'],
        [
            ['Multi-view (ego + exo)', 'More robust hand tracking', 'EgoExo4D-style cross-view learning; 3D reconstruction'],
            ['Multiple skill levels', 'More diverse training data', 'Skill assessment models (OSATS prediction); surgical education AI'],
            ['Audio narration', 'Language annotations for VITRA', 'Surgical narration generation; multimodal training; action-language grounding'],
            ['20 diverse tasks', 'Comprehensive action coverage', 'Surgical gesture taxonomy; transfer learning across procedures'],
            ['Calibrated cameras', 'Exact intrinsics for HaWoR', 'Any future hand tracker works; metric-scale reconstruction'],
            ['Multiple trials per task', 'Statistical reliability', 'Learning curve analysis; consistency metrics; temporal modeling'],
            ['Raw 30fps video', 'HaWoR temporal tracking', 'Optical flow; video prediction models; slow-motion analysis'],
            ['JIGSAWS-comparable tasks', 'Direct benchmarking', 'Cross-domain transfer (open surgery <-> laparoscopic)'],
            ['Structured metadata', 'Organized processing', 'Dataset filtering; cohort analysis; reproducible experiments'],
            ['Error annotations', 'Skill assessment', 'Surgical error detection; safety-critical AI research'],
        ]
    )

    # =================================================================
    # 13. COMPARISON
    # =================================================================
    doc.add_heading('13. Comparison with Existing Datasets', level=1)

    doc.add_paragraph(
        'The following table positions EgoSurg-Bench against all relevant datasets '
        'across the dimensions that matter for VLA fine-tuning:'
    )

    add_table(doc,
        ['', 'Ego video', 'Open surg.', 'MANO pose', 'Language', 'Multi-user', 'Calibrated', 'Public'],
        [
            ['EgoSurg-Bench (ours)', 'Yes (30fps)', 'Yes (pad)', 'Yes (HaWoR)', 'Yes (narration)', 'Yes (3-8)', 'Yes', 'Yes'],
            ['EgoSurgery', '0.5fps JPGs', 'Yes (real)', 'No', 'No', 'Yes (8)', 'No', 'Yes'],
            ['AVOS YouTube', 'No', 'Yes (real)', 'No', 'Action classes', 'Unknown', 'No', 'Yes'],
            ['POV-Surgery', 'Yes (synth)', 'Yes (synth)', 'Yes (GT)', 'No', 'No', 'Yes', 'Yes'],
            ['JIGSAWS', 'No', 'No (dVRK)', 'Robot kin.', 'Gestures', 'Yes (8)', 'Yes', 'Yes'],
            ['Ego4D', 'Yes', 'No', 'HaWoR', 'Narration', 'Yes (931)', 'Mixed', 'Yes'],
        ]
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('EgoSurg-Bench is the only dataset that checks all boxes for VITRA fine-tuning '
              'on open surgery.').bold = True
    p.add_run(
        ' While it uses training pads rather than real tissue, the hand motions are '
        'authentic (this is how surgeons train), the instruments are real, and the '
        'environment is realistic. Combined with EgoSurgery (real tissue, lower quality poses) '
        'and POV-Surgery (synthetic, MANO GT), it forms a complementary training mixture.'
    )

    # --- Save ---
    output_path = Path(__file__).resolve().parent.parent / 'doc' / 'EgoSurg-Bench_Data_Collection_Plan.docx'
    doc.save(str(output_path))
    print(f'Saved to: {output_path}')


if __name__ == '__main__':
    main()
